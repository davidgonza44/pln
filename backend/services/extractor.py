from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO


@dataclass
class ExtractionResult:
    filename: str
    extension: str
    text: str
    detail: str
    warnings: list[str]


class ExtractionError(Exception):
    pass


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_file(filename: str, content: bytes) -> ExtractionResult:
    extension = Path(filename).suffix.lower()
    warnings: list[str] = []

    if not content:
        raise ExtractionError("El archivo está vacío o no se pudo leer.")

    if extension == ".txt":
        text = _extract_txt(content)
        return ExtractionResult(filename, extension, text, "Texto plano leído localmente.", warnings)

    if extension == ".docx":
        text, detail = _extract_docx(content)
        return ExtractionResult(filename, extension, text, detail, warnings)

    if extension == ".xlsx":
        text, detail = _extract_xlsx(content)
        return ExtractionResult(filename, extension, text, detail, warnings)

    if extension == ".pdf":
        text, detail, pdf_warnings = _extract_pdf(content)
        warnings.extend(pdf_warnings)
        return ExtractionResult(filename, extension, text, detail, warnings)

    raise ExtractionError(
        f"Formato no soportado: {extension or 'sin extensión'}. Use .txt, .docx, .pdf textual o .xlsx."
    )


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return _clean_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise ExtractionError("No fue posible decodificar el TXT con UTF-8 ni Latin-1.")


def _extract_docx(content: bytes) -> tuple[str, str]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise ExtractionError("Falta instalar python-docx. Ejecute: pip install -r requirements.txt") from exc

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionError("No fue posible abrir el DOCX. Verifique que no esté protegido o corrupto.") from exc

    parts: list[str] = []
    paragraphs = 0
    table_cells = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs += 1
            style = getattr(p.style, "name", "") if p.style else ""
            if "Heading" in style or "Título" in style or "Titulo" in style:
                parts.append(f"\n{text}\n")
            else:
                parts.append(text)

    for table_idx, table in enumerate(doc.tables, start=1):
        parts.append(f"\n[TABLA {table_idx}]")
        for row in table.rows:
            values = []
            for cell in row.cells:
                cell_text = _clean_text(cell.text)
                values.append(cell_text)
                if cell_text:
                    table_cells += 1
            if any(values):
                parts.append(" | ".join(values))

    text = _clean_text("\n\n".join(parts))
    if not text:
        raise ExtractionError("El DOCX no contiene texto extraíble.")
    detail = f"DOCX extraído: {paragraphs} párrafos y {table_cells} celdas de tabla."
    return text, detail


def _extract_xlsx(content: bytes) -> tuple[str, str]:
    try:
        import openpyxl
    except Exception as exc:  # pragma: no cover
        raise ExtractionError("Falta instalar openpyxl. Ejecute: pip install -r requirements.txt") from exc

    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise ExtractionError("No fue posible abrir el XLSX. Verifique que no esté protegido o corrupto.") from exc

    parts: list[str] = []
    rows_used = 0
    for ws in wb.worksheets:
        parts.append(f"\n[HOJA: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            values = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(values):
                rows_used += 1
                parts.append(" | ".join(values))

    text = _clean_text("\n".join(parts))
    if not text:
        raise ExtractionError("El XLSX no contiene celdas con texto o valores extraíbles.")
    detail = f"XLSX extraído: {len(wb.worksheets)} hoja(s), {rows_used} fila(s) con contenido."
    return text, detail


def _extract_pdf(content: bytes) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise ExtractionError("Falta instalar pypdf. Ejecute: pip install -r requirements.txt") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionError("No fue posible abrir el PDF. Verifique que no esté protegido o corrupto.") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            raise ExtractionError("El PDF está cifrado/protegido y no pudo abrirse sin contraseña.")

    parts: list[str] = []
    pages_with_text = 0
    for idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = _clean_text(text)
        if text:
            pages_with_text += 1
            parts.append(f"\n[PÁGINA {idx}]\n{text}")

    final_text = _clean_text("\n\n".join(parts))
    if not final_text:
        warnings.append("El PDF no contiene texto embebido. Probablemente es escaneado y requiere OCR.")
        raise ExtractionError("No se extrajo texto del PDF. PDF escaneado requiere OCR, fuera del MVP actual.")

    if pages_with_text < len(reader.pages):
        warnings.append("Algunas páginas no tenían texto extraíble; podrían ser escaneadas o imágenes.")

    detail = f"PDF textual extraído: {pages_with_text}/{len(reader.pages)} página(s) con texto."
    return final_text, detail, warnings
